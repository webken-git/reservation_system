from django import forms
from django.utils.decorators import method_decorator
from django.views.decorators.cache import cache_page
from django.views.decorators.vary import vary_on_cookie
from rest_framework import viewsets, response, status, mixins
import datetime
import pytz
import os
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from users import permissions
from documents.models import DocumentTemplate, Document
from documents.serializers import DocumentTemplateSerializer, DocumentSerializer
from reservations.models import ApprovalApplication, DefferdPayment, UsageCategory, AgeCategory


# データの変更が頻繫にあるAPIのキャッシュの期限は5分
TIME_OUTS_5MINUTES = 60 * 5
# 1日
TIME_OUTS_1DAY = 60 * 60 * 24
# マスターデータのキャッシュの期限は30日
TIME_OUTS_1MONTH = TIME_OUTS_1DAY * 30


def insert_string(string, index, add_string):
  return string[:index] + add_string + string[index:]


def create_new_word(request):
  """
  必須フォームフィールド:
  id: Documents テーブル id
  approval_application: approval applications テーブル id
  任意:
  number: 書類の発行番号
  """
  BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
  now = datetime.datetime.now(pytz.timezone('Asia/Tokyo'))
  query = DocumentTemplate.objects.filter(id=request.data['id'])
  # DBから検索
  approval_applications = ApprovalApplication.objects.filter(id=request.data['approval_application_id'])
  usage_categorizes = UsageCategory.objects.filter(reservation__id=approval_applications[0].reservation.id)
  age_categorizes = AgeCategory.objects.filter(reservation__id=approval_applications[0].reservation.id)
  # 記入内容データ
  group_name = approval_applications[0].reservation.group_name
  reader_name = approval_applications[0].reservation.reader_name
  contact_name = approval_applications[0].reservation.contact_name
  address = approval_applications[0].reservation.address
  tel = approval_applications[0].reservation.tel
  approval = approval_applications[0].approval.name
  place = approval_applications[0].reservation.place.name
  usages = [usage_categorize.name for usage_categorize in usage_categorizes[0].usage.all()]
  ages = [age_categorize.name for age_categorize in age_categorizes[0].age.all()]
  start = approval_applications[0].reservation.start
  end = approval_applications[0].reservation.end
  start_hour = start.hour
  end_hour = end.hour
  purpose = approval_applications[0].reservation.purpose
  organizer_number = approval_applications[0].reservation.organizer_number
  participant_number = approval_applications[0].reservation.participant_number
  equipment = [e.name for e in approval_applications[0].reservation.equipment.all()]
  special_equipment = [sp.name for sp in approval_applications[0].reservation.special_equipment.all()]
  admission_fee = approval_applications[0].reservation.admission_fee
  conditions = approval_applications[0].conditions
  cancellation_reason = approval_applications[0].cancellation_reason
  created_at = approval_applications[0].reservation.created_at
  updated_at = approval_applications[0].updated_at
  usage_fee = approval_applications[0].usage_fee
  heating_fee = approval_applications[0].heating_fee
  electric_fee = approval_applications[0].electric_fee

  if query:
    file = BASE_DIR + query[0].url
    doc = docx.Document(file)
    doc.styles['Normal'].font.name = 'ＭＳ 明朝'
    tbl = doc.tables[0]

    if query[0].name == '稚内市体育施設使用等承認（不承認）通知書' or query[0].name == '稚内市体育施設使用等承認取消し等決定通知書' or query[0].name == '稚内市体育施設使用料等後納承認（不承認）通知書':
      number = request.data['number']  # 発行番号
      tbl.rows[0].cells[0].paragraphs[0].text = tbl.rows[0].cells[0].paragraphs[0].text.replace('第　　　　　号', '第　{}　号'.format(number))
      # 使用（利用）体育施設の名称
      tbl.rows[1].cells[1].paragraphs[0].text = insert_string(tbl.rows[1].cells[1].paragraphs[0].text, 5, place)
      # 使用（利用）区分
      if 'アマチュアスポーツ' in usages:
        tbl.rows[2].cells[1].paragraphs[0].text = tbl.rows[2].cells[1].paragraphs[0].text.replace('□アマチュア', '☑アマチュア')
      if '一般使用' in usages:
        tbl.rows[2].cells[1].paragraphs[0].text = tbl.rows[2].cells[1].paragraphs[0].text.replace('□一般使用', '☑一般使用')
      if '競技会使用' in usages:
        tbl.rows[2].cells[1].paragraphs[0].text = tbl.rows[2].cells[1].paragraphs[0].text.replace('□競技会使用', '☑競技会使用')
      if '非営利' in usages:
        tbl.rows[2].cells[1].paragraphs[1].text = tbl.rows[2].cells[1].paragraphs[1].text.replace('□非営利', '☑非営利')
        if '入場料を徴収する' in usages:
          tbl.rows[2].cells[1].paragraphs[1].text = tbl.rows[2].cells[1].paragraphs[1].text.replace('□入場料を徴収する', '☑入場料を徴収する')
        elif '入場料を徴収しない' in usages:
          tbl.rows[2].cells[1].paragraphs[1].text = tbl.rows[2].cells[1].paragraphs[1].text.replace('□入場料を徴収しない', '☑入場料を徴収しない')
      elif '営利' in usages:
        tbl.rows[2].cells[1].paragraphs[2].text = tbl.rows[2].cells[1].paragraphs[2].text.replace('□営　利', '☑営　利')
        if '入場料を徴収する' in usages:
          tbl.rows[2].cells[1].paragraphs[2].text = tbl.rows[2].cells[1].paragraphs[2].text.replace('□入場料を徴収する', '☑入場料を徴収する')
        elif '入場料を徴収しない' in usages:
          tbl.rows[2].cells[1].paragraphs[2].text = tbl.rows[2].cells[1].paragraphs[2].text.replace('□入場料を徴収しない', '☑入場料を徴収しない')
      # 年齢区分
      if '幼児' in ages:
        tbl.rows[3].cells[1].paragraphs[0].text = tbl.rows[3].cells[1].paragraphs[0].text.replace('□幼児', '☑幼児')
      if '小学生' in ages:
        tbl.rows[3].cells[1].paragraphs[0].text = tbl.rows[3].cells[1].paragraphs[0].text.replace('□小学生', '☑小学生')
      if '中学生' in ages:
        tbl.rows[3].cells[1].paragraphs[0].text = tbl.rows[3].cells[1].paragraphs[0].text.replace('□中学生', '☑中学生')
      if '高校生' in ages:
        tbl.rows[3].cells[1].paragraphs[0].text = tbl.rows[3].cells[1].paragraphs[0].text.replace('□高校生', '☑高校生')
      if '大学生' in ages:
        tbl.rows[3].cells[1].paragraphs[0].text = tbl.rows[3].cells[1].paragraphs[0].text.replace('□大学生', '☑大学生')
      if '一般' in ages:
        tbl.rows[3].cells[1].paragraphs[1].text = tbl.rows[3].cells[1].paragraphs[1].text.replace('□一般', '☑一般')
      if '高齢者' in ages:
        tbl.rows[3].cells[1].paragraphs[1].text = tbl.rows[3].cells[1].paragraphs[1].text.replace('□高齢者', '☑高齢者')
      if '障害者' in ages:
        tbl.rows[3].cells[1].paragraphs[1].text = tbl.rows[3].cells[1].paragraphs[1].text.replace('□障害者', '☑障害者')
      # 利用日時
      tbl.rows[4].cells[1].paragraphs[0].text = tbl.rows[4].cells[1].paragraphs[0].text.replace('年', str(start.year) + ' 年').replace('　月', str(start.month) + ' 月').replace('　日', str(start.day) + ' 日').replace('　時', str(start.strftime('%I')) + ' 時').replace('　分', str(start.minute) + ' 分')
      tbl.rows[4].cells[1].paragraphs[1].text = tbl.rows[4].cells[1].paragraphs[1].text.replace('年', str(end.year) + ' 年').replace('　月', str(end.month) + ' 月').replace('　日', str(end.day) + ' 日').replace('　時', str(end.strftime('%I')) + ' 時').replace('　分', str(end.minute) + ' 分')
      if start_hour - 12 < 0:
        tbl.rows[4].cells[1].paragraphs[0].text = tbl.rows[4].cells[1].paragraphs[0].text.replace('前', '☑前')
      elif start_hour - 12 > 0:
        tbl.rows[4].cells[1].paragraphs[0].text = tbl.rows[4].cells[1].paragraphs[0].text.replace('後', '☑後')
      if end_hour - 12 < 0:
        tbl.rows[4].cells[1].paragraphs[1].text = tbl.rows[4].cells[1].paragraphs[1].text.replace('前', '☑前')
      elif end_hour - 12 > 0:
        tbl.rows[4].cells[1].paragraphs[1].text = tbl.rows[4].cells[1].paragraphs[1].text.replace('後', '☑後')
    else:
      tbl.rows[0].cells[3].paragraphs[0].text = updated_at.strftime('%Y 年 %m 月 %d 日').replace('年 0', '年 ').replace('月 0', '月 ')
      # 右揃えにしてgroup_nameを挿入
      tbl.rows[0].cells[3].paragraphs[4].text = "　　　　　　　　　　　　　　　　　　　　　　　・団 体 名　{}".format(group_name)
      tbl.rows[0].cells[3].paragraphs[4].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
      tbl.rows[0].cells[3].paragraphs[5].text = "　　　　　　　　　　　　　　　　　　　　　　　　代表者名　{}".format(reader_name)
      # reader_nameの位置をgroup_nameの位置に合わせる
      tbl.rows[0].cells[3].paragraphs[5].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
      tbl.rows[0].cells[3].paragraphs[7].text = "　　　　　　　　　　　　　　　　　　　　　　　・連絡者名　{}".format(contact_name)
      tbl.rows[0].cells[3].paragraphs[7].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
      tbl.rows[0].cells[3].paragraphs[8].text = "　　　　　　　　　　　　　　　　　　　　　　　　住　　所　{}".format(address)
      tbl.rows[0].cells[3].paragraphs[8].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
      tbl.rows[0].cells[3].paragraphs[9].text = "　　　　　　　　　　　　　　　　　　　　　　　　電話番号　{}".format(tel)
      tbl.rows[0].cells[3].paragraphs[9].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
      # 使用（利用）体育施設の名称
      tbl.rows[1].cells[3].paragraphs[0].text = insert_string(tbl.rows[1].cells[3].paragraphs[0].text, 5, place)
      # 使用（利用）区分
      if 'アマチュアスポーツ' in usages:
        tbl.rows[2].cells[3].paragraphs[0].text = tbl.rows[2].cells[3].paragraphs[0].text.replace('□アマチュア', '☑アマチュア')
      if '一般使用' in usages:
        tbl.rows[2].cells[3].paragraphs[1].text = tbl.rows[2].cells[3].paragraphs[1].text.replace('□一般使用', '☑一般使用')
      if '競技会使用' in usages:
        tbl.rows[2].cells[3].paragraphs[1].text = tbl.rows[2].cells[3].paragraphs[1].text.replace('□競技会使用', '☑競技会使用')
      if '非営利' in usages:
        tbl.rows[2].cells[3].paragraphs[2].text = tbl.rows[2].cells[3].paragraphs[2].text.replace('□非営利', '☑非営利')
        if '入場料を徴収する' in usages:
          tbl.rows[2].cells[3].paragraphs[2].text = tbl.rows[2].cells[3].paragraphs[2].text.replace('□入場料を徴収する', '☑入場料を徴収する')
        elif '入場料を徴収しない' in usages:
          tbl.rows[2].cells[3].paragraphs[2].text = tbl.rows[2].cells[3].paragraphs[2].text.replace('□入場料を徴収しない', '☑入場料を徴収しない')
      elif '営利' in usages:
        tbl.rows[2].cells[3].paragraphs[3].text = tbl.rows[2].cells[3].paragraphs[3].text.replace('□営　利', '☑営　利')
        if '入場料を徴収する' in usages:
          tbl.rows[2].cells[3].paragraphs[3].text = tbl.rows[2].cells[3].paragraphs[3].text.replace('□入場料を徴収する', '☑入場料を徴収する')
        elif '入場料を徴収しない' in usages:
          tbl.rows[2].cells[3].paragraphs[3].text = tbl.rows[2].cells[3].paragraphs[3].text.replace('□入場料を徴収しない', '☑入場料を徴収しない')
      # 年齢区分
      if '幼児' in ages:
        tbl.rows[3].cells[3].paragraphs[0].text = tbl.rows[3].cells[3].paragraphs[0].text.replace('□幼児', '☑幼児')
      if '小学生' in ages:
        tbl.rows[3].cells[3].paragraphs[0].text = tbl.rows[3].cells[3].paragraphs[0].text.replace('□小学生', '☑小学生')
      if '中学生' in ages:
        tbl.rows[3].cells[3].paragraphs[0].text = tbl.rows[3].cells[3].paragraphs[0].text.replace('□中学生', '☑中学生')
      if '高校生' in ages:
        tbl.rows[3].cells[3].paragraphs[0].text = tbl.rows[3].cells[3].paragraphs[0].text.replace('□高校生', '☑高校生')
      if '大学生' in ages:
        tbl.rows[3].cells[3].paragraphs[0].text = tbl.rows[3].cells[3].paragraphs[0].text.replace('□大学生', '☑大学生')
      if '一般' in ages:
        tbl.rows[3].cells[3].paragraphs[1].text = tbl.rows[3].cells[3].paragraphs[1].text.replace('□一般', '☑一般')
      if '高齢者' in ages:
        tbl.rows[3].cells[3].paragraphs[1].text = tbl.rows[3].cells[3].paragraphs[1].text.replace('□高齢者', '☑高齢者')
      if '障害者' in ages:
        tbl.rows[3].cells[3].paragraphs[1].text = tbl.rows[3].cells[3].paragraphs[1].text.replace('□障害者', '☑障害者')
      # 利用日時
      tbl.rows[4].cells[3].paragraphs[0].text = tbl.rows[4].cells[3].paragraphs[0].text.replace('年', str(start.year) + ' 年').replace('　月', str(start.month) + ' 月').replace('　日', str(start.day) + ' 日').replace('　時', str(start.strftime('%I')) + ' 時').replace('　分', str(start.minute) + ' 分')
      tbl.rows[4].cells[3].paragraphs[1].text = tbl.rows[4].cells[3].paragraphs[1].text.replace('年', str(end.year) + ' 年').replace('　月', str(end.month) + ' 月').replace('　日', str(end.day) + ' 日').replace('　時', str(end.strftime('%I')) + ' 時').replace('　分', str(end.minute) + ' 分')
      if start_hour < 0:
        tbl.rows[4].cells[3].paragraphs[0].text = tbl.rows[4].cells[3].paragraphs[0].text.replace('前', '☑前')
      elif start_hour > 0:
        tbl.rows[4].cells[3].paragraphs[0].text = tbl.rows[4].cells[3].paragraphs[0].text.replace('後', '☑後')
      if end_hour < 0:
        tbl.rows[4].cells[3].paragraphs[1].text = tbl.rows[4].cells[3].paragraphs[1].text.replace('前', '☑前')
      elif end_hour > 0:
        tbl.rows[4].cells[3].paragraphs[1].text = tbl.rows[4].cells[3].paragraphs[1].text.replace('後', '☑後')

    # 稚内市体育施設使用等承認（不承認）通知書
    if query[0].name == '稚内市体育施設使用等承認（不承認）通知書':
      tbl.rows[0].cells[0].paragraphs[1].text = updated_at.strftime('%Y 年 %m 月 %d 日').replace('年 0', '年 ').replace('月 0', '月 ')
      tbl.rows[0].cells[0].paragraphs[3].text = insert_string(tbl.rows[0].cells[0].paragraphs[3].text, 6, "{}　{}".format(group_name, contact_name))
      tbl.rows[0].cells[0].paragraphs[10].text = tbl.rows[0].cells[0].paragraphs[10].text.replace('　　年', str(created_at.year) + ' 年').replace('　月', str(created_at.month) + ' 月').replace('　日', str(created_at.day) + ' 日')
      if approval == '承認':
        tbl.rows[0].cells[0].paragraphs[10].text = tbl.rows[0].cells[0].paragraphs[10].text.replace('□承認', '☑承認')
      elif approval == '不承認':
        tbl.rows[0].cells[0].paragraphs[10].text = tbl.rows[0].cells[0].paragraphs[10].text.replace('□不承認', '☑不承認')
      else:
        return {'error': '承認または不承認されたデータを指定してください。'}
      tbl.rows[5].cells[1].paragraphs[0].text = insert_string(tbl.rows[5].cells[1].paragraphs[0].text, 0, purpose)
      tbl.rows[6].cells[1].paragraphs[0].text = tbl.rows[6].cells[1].paragraphs[0].text.replace('者　　', '者　' + str(organizer_number)).replace('員　　', '員　' + str(participant_number))
      if equipment:
        tbl.rows[7].cells[1].paragraphs[0].text = insert_string(tbl.rows[7].cells[1].paragraphs[0].text, 3, str(equipment)).replace("'", '').replace('[', '').replace(']　　　　　　　　　　　　　　　　　　', '')
      else:
        tbl.rows[7].cells[1].paragraphs[0].text = tbl.rows[7].cells[1].paragraphs[0].text.replace('無', '✓無')
      if special_equipment:
        tbl.rows[8].cells[1].paragraphs[0].text = insert_string(tbl.rows[8].cells[1].paragraphs[0].text, 3, str(special_equipment)).replace("'", '').replace('[', '').replace(']　　　　　　　　　　　', '')
      else:
        tbl.rows[8].cells[1].paragraphs[0].text = tbl.rows[8].cells[1].paragraphs[0].text.replace('無', '✓無')
      if admission_fee:
        tbl.rows[9].cells[1].paragraphs[0].text = tbl.rows[9].cells[1].paragraphs[0].text.replace('円', str(admission_fee) + '円')
      else:
        tbl.rows[9].cells[1].paragraphs[0].text = tbl.rows[9].cells[1].paragraphs[0].text.replace('無', '✓無')
      if conditions:
        tbl.rows[13].cells[1].paragraphs[0].text = tbl.rows[13].cells[1].paragraphs[0].text.replace('）', '）\n' + str(conditions))
    # 稚内市体育施設使用等承認取消し等決定通知書
    elif query[0].name == '稚内市体育施設使用等承認取消し等決定通知書':
      if cancellation_reason:
        tbl.rows[0].cells[0].paragraphs[1].text = updated_at.strftime('%Y 年 %m 月 %d 日').replace('年 0', '年 ').replace('月 0', '月 ')
        tbl.rows[0].cells[0].paragraphs[2].text = insert_string(tbl.rows[0].cells[0].paragraphs[2].text, 6, "{}　{}".format(group_name, contact_name))
        tbl.rows[0].cells[0].paragraphs[9].text = tbl.rows[0].cells[0].paragraphs[9].text.replace('　　年', str(created_at.year) + ' 年').replace('　月', str(created_at.month) + ' 月').replace('　日', str(created_at.day) + ' 日').replace('第　　　　号', '第　' + str(number) + '　号')
        tbl.rows[5].cells[1].paragraphs[0].text = insert_string(tbl.rows[5].cells[1].paragraphs[0].text, 0, purpose)
        tbl.rows[6].cells[0].paragraphs[2].text = insert_string(tbl.rows[6].cells[0].paragraphs[2].text, 0, str(cancellation_reason))
      else:
        return {'error': 'approval-applicationテーブルにあるcancellation_reasonフィールドの値が未入力です。'}
    # 稚内市体育施設使用等承認申請書
    elif query[0].name == '稚内市体育施設使用等承認申請書':
      tbl.rows[5].cells[3].paragraphs[0].text = insert_string(tbl.rows[5].cells[3].paragraphs[0].text, 0, purpose)
      tbl.rows[6].cells[3].paragraphs[0].text = tbl.rows[6].cells[3].paragraphs[0].text.replace('者　　', '者　' + str(organizer_number)).replace('員　　', '員　' + str(participant_number))
      if equipment:
        tbl.rows[7].cells[3].paragraphs[0].text = insert_string(tbl.rows[7].cells[3].paragraphs[0].text, 3, str(equipment)).replace("'", '').replace('[', '').replace(']　　　　　　　　　　　　　　　　　　', '')
      else:
        tbl.rows[9].cells[3].paragraphs[0].text = tbl.rows[9].cells[3].paragraphs[0].text.replace('無', '✓無')
      if special_equipment:
        tbl.rows[8].cells[3].paragraphs[0].text = insert_string(tbl.rows[8].cells[3].paragraphs[0].text, 3, str(special_equipment)).replace("'", '').replace('[', '').replace(']　　　　　　　　　　　', '')
      else:
        tbl.rows[9].cells[3].paragraphs[0].text = tbl.rows[9].cells[3].paragraphs[0].text.replace('無', '✓無')
      if admission_fee:
        tbl.rows[9].cells[3].paragraphs[0].text = tbl.rows[9].cells[3].paragraphs[0].text.replace('円', str(admission_fee) + '円')
      else:
        tbl.rows[9].cells[3].paragraphs[0].text = tbl.rows[9].cells[3].paragraphs[0].text.replace('無', '✓無')
      if conditions:
        tbl.rows[16].cells[1].paragraphs[0].text = tbl.rows[16].cells[1].paragraphs[0].text.replace('）', '）\n' + str(conditions))
    # 稚内市体育施設使用料等後納承認（不承認）通知書
    elif query[0].name == '稚内市体育施設使用料等後納承認（不承認）通知書':
      q = DefferdPayment.objects.filter(reservation=approval_applications[0].reservation.id)
      tbl.rows[6].cells[1].paragraphs[1].text = insert_string(tbl.rows[6].cells[1].paragraphs[1].text.replace('（　　　　　　　　　　　　　　　　　　　　　　　　　　', '（'), 2, q[0].reason)
      if conditions:
        tbl.rows[7].cells[1].paragraphs[0].text = tbl.rows[7].cells[1].paragraphs[0].text.replace('）', '）\n' + str(conditions))
    # 稚内市体育施設使用料等後納申請書
    elif query[0].name == '稚内市体育施設使用料等後納申請書':
      q = DefferdPayment.objects.filter(reservation=approval_applications[0].reservation.id)
      tbl.rows[6].cells[3].paragraphs[1].text = insert_string(tbl.rows[6].cells[3].paragraphs[1].text.replace('（　　　　　　　　　　　　　　　　　　　　　　　　　　', '（'), 2, q[0].reason)
      if conditions:
        tbl.rows[10].cells[3].paragraphs[0].text = tbl.rows[10].cells[3].paragraphs[0].text.replace('）', '）\n' + str(conditions))
  else:
    return {'error': 'docxファイルの指定が違います。'}
  doc.save(BASE_DIR + '/static/documents/docx/' + now.strftime('%Y%m%d-%H%M%S_') + str(query[0].id) + '.docx')
  return now.strftime('%Y%m%d-%H%M%S_') + str(query[0].id) + '.docx', now.strftime('%Y%m%d-%H%M%S_') + query[0].name + '.docx'


class DocumentTemplateViewSet(viewsets.ModelViewSet):
  queryset = DocumentTemplate.objects.all()
  serializer_class = DocumentTemplateSerializer
  filter_fields = [f.name for f in DocumentTemplate._meta.fields]
  # permission_classes = [permissions.ActionBasedPermission]
  action_permissions = {
      permissions.IsAdminUser: ['update', 'partial_update', 'create', 'destroy'],
      permissions.IsAuthenticated: ['list', 'retrieve'],
      permissions.AllowAny: []
  }

  def create(self, request, *args, **kwargs):
    serializer = self.get_serializer(data=request.data)
    serializer.is_valid(raise_exception=True)
    self.perform_create(serializer)
    headers = self.get_success_headers(serializer.data)
    return response.Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)

  @method_decorator(vary_on_cookie)
  @method_decorator(cache_page(TIME_OUTS_5MINUTES))
  def list(self, request, *args, **kwargs):
    return super().list(request, *args, **kwargs)

  # @method_decorator(vary_on_cookie)
  # @method_decorator(cache_page(TIME_OUTS_5MINUTES))
  def retrieve(self, request, *args, **kwargs):
    return super().retrieve(request, *args, **kwargs)


class DocumentViewSet(viewsets.ModelViewSet):
  """
  申請書の作成と削除
  """
  queryset = Document.objects.all()
  serializer_class = DocumentSerializer
  # permission_classes = [permissions.ActionBasedPermission]
  action_permissions = {
      permissions.IsAdminUser: ['list', 'retrieve', 'update', 'partial_update'],
      permissions.IsAuthenticated: [],
      permissions.AllowAny: ['create', 'destroy', ]
  }

  def create(self, request, *args, **kwargs):
    file_name = create_new_word(self.request)
    if 'error' in file_name:
      return response.Response(file_name, status=status.HTTP_400_BAD_REQUEST)
    else:
      serializer = self.get_serializer(data=request.data)
      serializer.is_valid(raise_exception=True)
      self.perform_create(serializer)
      headers = self.get_success_headers(serializer.data)
      return response.Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)

  def perform_create(self, serializer):
    file, file_name = create_new_word(self.request)
    serializer.save(
        number=self.request.data['number'],
        file=file,
        file_name=file_name,
        approval_application_id=self.request.data['approval_application_id']
    )
    return response.Response(serializer.data, status=status.HTTP_201_CREATED)

  def destroy(self, request, *args, **kwargs):
    instance = self.get_object()
    serializer = self.get_serializer(instance)
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    if os.path.exists(BASE_DIR + '/static/application_documents/docx/' + instance.file):
      os.remove(BASE_DIR + '/static/application_documents/docx/' + instance.file)
    self.perform_destroy(instance)
    return response.Response(serializer.data, status=status.HTTP_200_OK)

  def perform_destroy(self, instance):
    instance.delete()
