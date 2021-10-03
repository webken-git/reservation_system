from django.utils.decorators import method_decorator
from django.views.decorators.cache import cache_page
from django.views.decorators.vary import vary_on_cookie
from rest_framework import viewsets, response, status, mixins
import datetime
import pytz
import os
import docx
from users import permissions
from application_documents.models import *
from application_documents.serializers import DocumentSerializer
from reservations.models import ApprovalApplication, UsageCategory, AgeCategory


# データの変更が頻繫にあるAPIのキャッシュの期限は5分
TIME_OUTS_5MINUTES = 60 * 5
# 1日
TIME_OUTS_1DAY = 60 * 60 * 24
# マスターデータのキャッシュの期限は30日
TIME_OUTS_1MONTH = TIME_OUTS_1DAY * 30


def insert_string(string, index, add_string):
  return string[:index] + add_string + string[index:]


def create_new_word(request):
  """
  必須フォームフィールド:
  id: Documents テーブル id
  approvalapplication_id: approval applications テーブル id
  任意:
  number: 書類の発行番号
  """
  BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
  now = datetime.datetime.now(pytz.timezone('Asia/Tokyo'))
  query = Document.objects.filter(id=request.data['id'])
  # DBから検索
  approval_applications = ApprovalApplication.objects.filter(id=request.data['approvalapplication_id'])
  usage_categorizes = UsageCategory.objects.filter(reservation__id=approval_applications[0].reservation.id)
  age_categorizes = AgeCategory.objects.filter(reservation__id=approval_applications[0].reservation.id)
  # 記入内容データ
  contact_name = approval_applications[0].reservation.contact_name
  address = approval_applications[0].reservation.address
  tel = approval_applications[0].reservation.tel
  approval = approval_applications[0].approval.name
  place = approval_applications[0].reservation.place.name
  usages = [usage_categorize.name for usage_categorize in usage_categorizes[0].usage.all()]
  ages = [age_categorize.name for age_categorize in age_categorizes[0].age.all()]
  start = approval_applications[0].reservation.start
  end = approval_applications[0].reservation.end
  start_hour = start.hour
  end_hour = end.hour
  purpose = approval_applications[0].reservation.purpose
  organizer_number = approval_applications[0].reservation.organizer_number
  participant_number = approval_applications[0].reservation.participant_number
  equipment = [e.name for e in approval_applications[0].reservation.equipment.all()]
  special_equipment = [sp.name for sp in approval_applications[0].reservation.special_equipment.all()]
  admission_fee = approval_applications[0].reservation.admission_fee
  conditions = approval_applications[0].conditions
  cancellation_reason = approval_applications[0].cancellation_reason

  if query:
    file = BASE_DIR + query[0].url
    doc = docx.Document(file)
    tbl = doc.tables[0]

    if request.data['id'] == '1' or request.data['id'] == '2' or request.data['id'] == '4':
      number = request.data['number']  # 発行番号
      # 使用（利用）体育施設の名称
      tbl.rows[1].cells[1].paragraphs[0].text = insert_string(tbl.rows[1].cells[1].paragraphs[0].text, 5, place)
      # 使用（利用）区分
      if 'アマチュアスポーツ' in usages:
        tbl.rows[2].cells[1].paragraphs[0].text = tbl.rows[2].cells[1].paragraphs[0].text.replace('□アマチュア', '☒アマチュア')
      if '一般使用' in usages:
        tbl.rows[2].cells[1].paragraphs[0].text = tbl.rows[2].cells[1].paragraphs[0].text.replace('□一般使用', '☒一般使用')
      if '競技会使用' in usages:
        tbl.rows[2].cells[1].paragraphs[0].text = tbl.rows[2].cells[1].paragraphs[0].text.replace('□競技会使用', '☒競技会使用')
      if '非営利' in usages:
        tbl.rows[2].cells[1].paragraphs[1].text = tbl.rows[2].cells[1].paragraphs[1].text.replace('□非営利', '☒非営利')
        if '入場料を徴収する' in usages:
          tbl.rows[2].cells[1].paragraphs[1].text = tbl.rows[2].cells[1].paragraphs[1].text.replace('□入場料を徴収する', '☒入場料を徴収する')
        elif '入場料を徴収しない' in usages:
          tbl.rows[2].cells[1].paragraphs[1].text = tbl.rows[2].cells[1].paragraphs[1].text.replace('□入場料を徴収しない', '☒入場料を徴収しない')
      elif '営利' in usages:
        tbl.rows[2].cells[1].paragraphs[2].text = tbl.rows[2].cells[1].paragraphs[2].text.replace('□営　利', '☒営　利')
        if '入場料を徴収する' in usages:
          tbl.rows[2].cells[1].paragraphs[2].text = tbl.rows[2].cells[1].paragraphs[2].text.replace('□入場料を徴収する', '☒入場料を徴収する')
        elif '入場料を徴収しない' in usages:
          tbl.rows[2].cells[1].paragraphs[2].text = tbl.rows[2].cells[1].paragraphs[2].text.replace('□入場料を徴収しない', '☒入場料を徴収しない')
      # 年齢区分
      if '幼児' in ages:
        tbl.rows[3].cells[1].paragraphs[0].text = tbl.rows[3].cells[1].paragraphs[0].text.replace('□幼児', '☒幼児')
      if '小学生' in ages:
        tbl.rows[3].cells[1].paragraphs[0].text = tbl.rows[3].cells[1].paragraphs[0].text.replace('□小学生', '☒小学生')
      if '中学生' in ages:
        tbl.rows[3].cells[1].paragraphs[0].text = tbl.rows[3].cells[1].paragraphs[0].text.replace('□中学生', '☒中学生')
      if '高校生' in ages:
        tbl.rows[3].cells[1].paragraphs[0].text = tbl.rows[3].cells[1].paragraphs[0].text.replace('□高校生', '☒高校生')
      if '大学生' in ages:
        tbl.rows[3].cells[1].paragraphs[0].text = tbl.rows[3].cells[1].paragraphs[0].text.replace('□大学生', '☒大学生')
      if '一般' in ages:
        tbl.rows[3].cells[1].paragraphs[1].text = tbl.rows[3].cells[1].paragraphs[1].text.replace('□一般', '☒一般')
      if '高齢者' in ages:
        tbl.rows[3].cells[1].paragraphs[1].text = tbl.rows[3].cells[1].paragraphs[1].text.replace('□高齢者', '☒高齢者')
      if '障害者' in ages:
        tbl.rows[3].cells[1].paragraphs[1].text = tbl.rows[3].cells[1].paragraphs[1].text.replace('□障害者', '☒障害者')
      # 利用日時
      tbl.rows[4].cells[1].paragraphs[0].text = tbl.rows[4].cells[1].paragraphs[0].text.replace('年', str(start.year) + ' 年').replace('　月', str(start.month) + ' 月').replace('　日', str(start.day) + ' 日').replace('　時', str(start.strftime('%I')) + ' 時').replace('　分', str(start.minute) + ' 分')
      tbl.rows[4].cells[1].paragraphs[1].text = tbl.rows[4].cells[1].paragraphs[1].text.replace('年', str(end.year) + ' 年').replace('　月', str(end.month) + ' 月').replace('　日', str(end.day) + ' 日').replace('　時', str(end.strftime('%I')) + ' 時').replace('　分', str(end.minute) + ' 分')
      if start_hour - 12 < 0:
        tbl.rows[4].cells[1].paragraphs[0].text = tbl.rows[4].cells[1].paragraphs[0].text.replace('前', '☒前')
      elif start_hour - 12 > 0:
        tbl.rows[4].cells[1].paragraphs[0].text = tbl.rows[4].cells[1].paragraphs[0].text.replace('後', '☒後')
      if end_hour - 12 < 0:
        tbl.rows[4].cells[1].paragraphs[1].text = tbl.rows[4].cells[1].paragraphs[1].text.replace('前', '☒前')
      elif end_hour - 12 > 0:
        tbl.rows[4].cells[1].paragraphs[1].text = tbl.rows[4].cells[1].paragraphs[1].text.replace('後', '☒後')
    else:
      tbl.rows[0].cells[3].paragraphs[0].text = now.strftime('%Y 年 %m 月 %d 日').replace('年 0', '年 ').replace('月 0', '月 ')
      if approval_applications[0].reservation.is_group == 1:
        tbl.rows[0].cells[3].paragraphs[4].text = tbl.rows[0].cells[3].paragraphs[4].text.replace('名　　　　', '名　' + approval_applications[0].reservation.group_name + ' ')
      elif approval_applications[0].reservation.is_group == 0:
        tbl.rows[0].cells[3].paragraphs[5].text = tbl.rows[0].cells[3].paragraphs[5].text.replace('名　　　　', '名　' + approval_applications[0].reservation.reader_name + ' ')
      tbl.rows[0].cells[3].paragraphs[7].text = tbl.rows[0].cells[3].paragraphs[7].text.replace('名　　　　', '名　' + contact_name + ' ')
      tbl.rows[0].cells[3].paragraphs[8].text = tbl.rows[0].cells[3].paragraphs[8].text.replace('所　　　　　　　　　　', '所　' + address + '')
      tbl.rows[0].cells[3].paragraphs[9].text = tbl.rows[0].cells[3].paragraphs[9].text.replace('号　　　　　　', '号　' + str(tel))
      # 使用（利用）体育施設の名称
      tbl.rows[1].cells[3].paragraphs[0].text = insert_string(tbl.rows[1].cells[3].paragraphs[0].text, 5, place)
      # 使用（利用）区分
      if 'アマチュアスポーツ' in usages:
        tbl.rows[2].cells[3].paragraphs[0].text = tbl.rows[2].cells[3].paragraphs[0].text.replace('□アマチュア', '☒アマチュア')
      if '一般使用' in usages:
        tbl.rows[2].cells[3].paragraphs[1].text = tbl.rows[2].cells[3].paragraphs[1].text.replace('□一般使用', '☒一般使用')
      if '競技会使用' in usages:
        tbl.rows[2].cells[3].paragraphs[1].text = tbl.rows[2].cells[3].paragraphs[1].text.replace('□競技会使用', '☒競技会使用')
      if '非営利' in usages:
        tbl.rows[2].cells[3].paragraphs[2].text = tbl.rows[2].cells[3].paragraphs[2].text.replace('□非営利', '☒非営利')
        if '入場料を徴収する' in usages:
          tbl.rows[2].cells[3].paragraphs[2].text = tbl.rows[2].cells[3].paragraphs[2].text.replace('□入場料を徴収する', '☒入場料を徴収する')
        elif '入場料を徴収しない' in usages:
          tbl.rows[2].cells[3].paragraphs[2].text = tbl.rows[2].cells[3].paragraphs[2].text.replace('□入場料を徴収しない', '☒入場料を徴収しない')
      elif '営利' in usages:
        tbl.rows[2].cells[3].paragraphs[3].text = tbl.rows[2].cells[3].paragraphs[3].text.replace('□営　利', '☒営　利')
        if '入場料を徴収する' in usages:
          tbl.rows[2].cells[3].paragraphs[3].text = tbl.rows[2].cells[3].paragraphs[3].text.replace('□入場料を徴収する', '☒入場料を徴収する')
        elif '入場料を徴収しない' in usages:
          tbl.rows[2].cells[3].paragraphs[3].text = tbl.rows[2].cells[3].paragraphs[3].text.replace('□入場料を徴収しない', '☒入場料を徴収しない')
      # 年齢区分
      if '幼児' in ages:
        tbl.rows[3].cells[3].paragraphs[0].text = tbl.rows[3].cells[3].paragraphs[0].text.replace('□幼児', '☒幼児')
      if '小学生' in ages:
        tbl.rows[3].cells[3].paragraphs[0].text = tbl.rows[3].cells[3].paragraphs[0].text.replace('□小学生', '☒小学生')
      if '中学生' in ages:
        tbl.rows[3].cells[3].paragraphs[0].text = tbl.rows[3].cells[3].paragraphs[0].text.replace('□中学生', '☒中学生')
      if '高校生' in ages:
        tbl.rows[3].cells[3].paragraphs[0].text = tbl.rows[3].cells[3].paragraphs[0].text.replace('□高校生', '☒高校生')
      if '大学生' in ages:
        tbl.rows[3].cells[3].paragraphs[0].text = tbl.rows[3].cells[3].paragraphs[0].text.replace('□大学生', '☒大学生')
      if '一般' in ages:
        tbl.rows[3].cells[3].paragraphs[1].text = tbl.rows[3].cells[3].paragraphs[1].text.replace('□一般', '☒一般')
      if '高齢者' in ages:
        tbl.rows[3].cells[3].paragraphs[1].text = tbl.rows[3].cells[3].paragraphs[1].text.replace('□高齢者', '☒高齢者')
      if '障害者' in ages:
        tbl.rows[3].cells[3].paragraphs[1].text = tbl.rows[3].cells[3].paragraphs[1].text.replace('□障害者', '☒障害者')
      # 利用日時
      tbl.rows[4].cells[3].paragraphs[0].text = tbl.rows[4].cells[3].paragraphs[0].text.replace('年', str(start.year) + ' 年').replace('　月', str(start.month) + ' 月').replace('　日', str(start.day) + ' 日').replace('　時', str(start.strftime('%I')) + ' 時').replace('　分', str(start.minute) + ' 分')
      tbl.rows[4].cells[3].paragraphs[1].text = tbl.rows[4].cells[3].paragraphs[1].text.replace('年', str(end.year) + ' 年').replace('　月', str(end.month) + ' 月').replace('　日', str(end.day) + ' 日').replace('　時', str(end.strftime('%I')) + ' 時').replace('　分', str(end.minute) + ' 分')
      if start_hour < 0:
        tbl.rows[4].cells[3].paragraphs[0].text = tbl.rows[4].cells[3].paragraphs[0].text.replace('前', '☒前')
      elif start_hour > 0:
        tbl.rows[4].cells[3].paragraphs[0].text = tbl.rows[4].cells[3].paragraphs[0].text.replace('後', '☒後')
      if end_hour < 0:
        tbl.rows[4].cells[3].paragraphs[1].text = tbl.rows[4].cells[3].paragraphs[1].text.replace('前', '☒前')
      elif end_hour > 0:
        tbl.rows[4].cells[3].paragraphs[1].text = tbl.rows[4].cells[3].paragraphs[1].text.replace('後', '☒後')

    # 稚内市体育施設使用等承認（不承認）通知書
    if query[0].name == '稚内市体育施設使用等承認（不承認）通知書':
      tbl.rows[0].cells[0].paragraphs[0].text = insert_string(tbl.rows[0].cells[0].paragraphs[0].text, 4, str(number))
      tbl.rows[0].cells[0].paragraphs[1].text = now.strftime('%Y 年 %m 月 %d 日').replace('年 0', '年 ').replace('月 0', '月 ')
      tbl.rows[0].cells[0].paragraphs[3].text = insert_string(tbl.rows[0].cells[0].paragraphs[3].text, 6, contact_name)
      tbl.rows[0].cells[0].paragraphs[10].text = tbl.rows[0].cells[0].paragraphs[10].text.replace('　　年', str(now.year) + ' 年').replace('　月', str(now.month) + ' 月').replace('　日', str(now.day) + ' 日')
      if approval == '承認':
        tbl.rows[0].cells[0].paragraphs[10].text = tbl.rows[0].cells[0].paragraphs[10].text.replace('□承認', '☒承認')
      elif approval == '不承認':
        tbl.rows[0].cells[0].paragraphs[10].text = tbl.rows[0].cells[0].paragraphs[10].text.replace('□不承認', '☒不承認')
      else:
        return {'error': '承認または不承認されたデータを指定してください。'}
      tbl.rows[5].cells[1].paragraphs[0].text = insert_string(tbl.rows[5].cells[1].paragraphs[0].text, 0, purpose)
      tbl.rows[6].cells[1].paragraphs[0].text = tbl.rows[6].cells[1].paragraphs[0].text.replace('者　　', '者　' + str(organizer_number)).replace('員　　', '員　' + str(participant_number))
      if equipment:
        tbl.rows[7].cells[1].paragraphs[0].text = insert_string(tbl.rows[7].cells[1].paragraphs[0].text, 3, str(equipment)).replace("'", '').replace('[', '').replace(']　　　　　　　　　　　　　　　　　　', '')
      else:
        tbl.rows[9].cells[1].paragraphs[0].text = tbl.rows[9].cells[1].paragraphs[0].text.replace('無', '✓無')
      if special_equipment:
        tbl.rows[8].cells[1].paragraphs[0].text = insert_string(tbl.rows[8].cells[1].paragraphs[0].text, 3, str(special_equipment)).replace("'", '').replace('[', '').replace(']　　　　　　　　　　　', '')
      else:
        tbl.rows[9].cells[1].paragraphs[0].text = tbl.rows[9].cells[1].paragraphs[0].text.replace('無', '✓無')
      if admission_fee:
        tbl.rows[9].cells[1].paragraphs[0].text = tbl.rows[9].cells[1].paragraphs[0].text.replace('円', str(admission_fee) + '円')
      else:
        tbl.rows[9].cells[1].paragraphs[0].text = tbl.rows[9].cells[1].paragraphs[0].text.replace('無', '✓無')
    elif query[0].name == '稚内市体育施設使用等承認取消し等決定通知書':
      tbl.rows[0].cells[0].paragraphs[0].text = insert_string(tbl.rows[0].cells[0].paragraphs[0].text, 4, str(number))
      tbl.rows[0].cells[0].paragraphs[1].text = now.strftime('%Y 年 %m 月 %d 日').replace('年 0', '年 ').replace('月 0', '月 ')
      tbl.rows[0].cells[0].paragraphs[2].text = insert_string(tbl.rows[0].cells[0].paragraphs[2].text, 6, contact_name)
      tbl.rows[0].cells[0].paragraphs[9].text = tbl.rows[0].cells[0].paragraphs[9].text.replace('　　年', str(now.year) + ' 年').replace('　月', str(now.month) + ' 月').replace('　日', str(now.day) + ' 日').replace('第　　　　号', '第　' + str(number) + '　号')
      tbl.rows[5].cells[1].paragraphs[0].text = insert_string(tbl.rows[5].cells[1].paragraphs[0].text, 0, purpose)
      tbl.rows[6].cells[0].paragraphs[2].text = insert_string(tbl.rows[6].cells[0].paragraphs[2].text, 0, cancellation_reason)
    elif query[0].name == '稚内市体育施設使用等承認申請書':
      tbl.rows[5].cells[3].paragraphs[0].text = insert_string(tbl.rows[5].cells[3].paragraphs[0].text, 0, purpose)
      tbl.rows[6].cells[3].paragraphs[0].text = tbl.rows[6].cells[3].paragraphs[0].text.replace('者　　', '者　' + str(organizer_number)).replace('員　　', '員　' + str(participant_number))
      if equipment:
        tbl.rows[7].cells[3].paragraphs[0].text = insert_string(tbl.rows[7].cells[3].paragraphs[0].text, 3, str(equipment)).replace("'", '').replace('[', '').replace(']　　　　　　　　　　　　　　　　　　', '')
      else:
        tbl.rows[9].cells[3].paragraphs[0].text = tbl.rows[9].cells[3].paragraphs[0].text.replace('無', '✓無')
      if special_equipment:
        tbl.rows[8].cells[3].paragraphs[0].text = insert_string(tbl.rows[8].cells[3].paragraphs[0].text, 3, str(special_equipment)).replace("'", '').replace('[', '').replace(']　　　　　　　　　　　', '')
      else:
        tbl.rows[9].cells[3].paragraphs[0].text = tbl.rows[9].cells[3].paragraphs[0].text.replace('無', '✓無')
      if admission_fee:
        tbl.rows[9].cells[3].paragraphs[0].text = tbl.rows[9].cells[3].paragraphs[0].text.replace('円', str(admission_fee) + '円')
      else:
        tbl.rows[9].cells[3].paragraphs[0].text = tbl.rows[9].cells[3].paragraphs[0].text.replace('無', '✓無')
      tbl.rows[16].cells[1].paragraphs[0].text = tbl.rows[16].cells[1].paragraphs[0].text.replace('）', '）\n' + str(conditions))
    elif query[0].name == '稚内市体育施設使用料等後納承認（不承認）通知書':
      pass
    elif query[0].name == '稚内市体育施設使用料等後納申請書':
      pass
  else:
    return {'error': 'docxファイルの指定が違います。'}
  doc.save(BASE_DIR + '/static/application_documents/docx/' + now.strftime('%Y%m%d_') + query[0].name + '.docx')
  # doc.save(BASE_DIR + '/static/application_documents/docx/' + now.strftime('%Y%m%d-%H%M%S_') + query[0].name + '.docx')
  return '/backend/django/static/application_documents/docx/' + now.strftime('%Y%m%d-%H%M%S_') + query[0].name + '.docx'
  # return {'path': tbl.rows[16].cells[0].paragraphs[0].text}


class DocumentViewSet(viewsets.ModelViewSet):
  queryset = Document.objects.all()
  serializer_class = DocumentSerializer
  filter_fields = [f.name for f in Document._meta.fields]
  permission_classes = [permissions.ActionBasedPermission]
  action_permissions = {
      permissions.IsAdminUser: ['update', 'partial_update', 'create', 'destroy'],
      permissions.IsAuthenticated: ['list', 'retrieve'],
      permissions.AllowAny: []
  }

  def create(self, request, *args, **kwargs):
    serializer = self.get_serializer(data=request.data)
    serializer.is_valid(raise_exception=True)
    self.perform_create(serializer)
    headers = self.get_success_headers(serializer.data)
    return response.Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)

  @ method_decorator(vary_on_cookie)
  @ method_decorator(cache_page(TIME_OUTS_5MINUTES))
  def list(self, request, *args, **kwargs):
    return super().list(request, *args, **kwargs)

  @ method_decorator(vary_on_cookie)
  @ method_decorator(cache_page(TIME_OUTS_5MINUTES))
  def retrieve(self, request, *args, **kwargs):
    return super().retrieve(request, *args, **kwargs)


class CreateNewDocumentViewSet(
        mixins.CreateModelMixin,
        mixins.DestroyModelMixin,
        viewsets.GenericViewSet):
  """
  申請書の作成と削除
  """
  queryset = Document.objects.all()
  serializer_class = DocumentSerializer
  permission_classes = [permissions.ActionBasedPermission]
  action_permissions = {
      permissions.IsAdminUser: [],
      permissions.IsAuthenticated: ['destroy'],
      permissions.AllowAny: ['create']
  }

  def create(self, request, *args, **kwargs):
    new_document = create_new_word(request)
    if new_document:
      return response.Response(new_document, status=status.HTTP_200_OK)
    else:
      return response.Response({'message': '失敗しました。'}, status=status.HTTP_400_BAD_REQUEST)

  def destroy(self, request, *args, **kwargs):
    return response.Response({'message': '削除しました。'}, status=status.HTTP_200_OK)
